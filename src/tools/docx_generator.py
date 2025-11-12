"""DOCX report generator for risk assessment reports.

This module creates professional Word documents with charts, tables,
and formatted text for risk assessment deliverables.
"""

import os
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path
import logging

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Chart generation
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import io

from ..models.schemas import (
    RiskAssessmentReport,
    ExecutiveSummary,
    RiskRating,
    VulnerabilityAnalysis,
)

logger = logging.getLogger(__name__)


class DOCXGenerator:
    """Generator for professional risk assessment reports in DOCX format.

    Creates reports with:
    - Executive summary
    - Risk heatmap
    - Vulnerability findings table
    - Threat intelligence summary
    - Risk ratings with justifications
    - Recommendations
    """

    # Color scheme
    COLORS = {
        "critical": RGBColor(220, 20, 60),  # Crimson
        "high": RGBColor(255, 99, 71),  # Tomato
        "medium": RGBColor(255, 165, 0),  # Orange
        "low": RGBColor(34, 139, 34),  # Forest Green
        "header": RGBColor(31, 78, 120),  # Dark Blue
    }

    def __init__(self, output_dir: str = "reports"):
        """Initialize DOCX generator.

        Args:
            output_dir: Directory for saving reports
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"DOCX generator initialized. Output dir: {output_dir}")

    def create_report(
        self,
        report_data: RiskAssessmentReport,
        filename: Optional[str] = None,
    ) -> str:
        """Create complete risk assessment report.

        Args:
            report_data: RiskAssessmentReport object
            filename: Output filename (auto-generated if None)

        Returns:
            Path to generated report

        Example:
            >>> generator = DOCXGenerator()
            >>> report_path = generator.create_report(report_data)
            >>> print(f"Report saved to: {report_path}")
        """
        logger.info("Creating risk assessment report")

        # Create document
        doc = Document()

        # Set document properties
        self._set_document_properties(doc, report_data)

        # Add title page
        self._add_title_page(doc, report_data)

        # Add executive summary
        self._add_executive_summary(doc, report_data.executive_summary)

        # Add risk heatmap
        self._add_risk_heatmap(doc, report_data.risk_ratings)

        # Add findings overview
        self._add_findings_overview(doc, report_data)

        # Add detailed vulnerabilities
        self._add_vulnerability_details(doc, report_data.vulnerabilities)

        # Add threat intelligence
        self._add_threat_intelligence(doc, report_data.threats)

        # Add risk ratings
        self._add_risk_ratings(doc, report_data.risk_ratings)

        # Add recommendations
        self._add_recommendations(doc, report_data)

        # Add appendices
        self._add_appendices(doc, report_data)

        # Generate filename
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"risk_assessment_report_{timestamp}.docx"

        output_path = os.path.join(self.output_dir, filename)

        # Save document
        doc.save(output_path)
        logger.info(f"Report saved to: {output_path}")

        return output_path

    def _set_document_properties(
        self, doc: Document, report_data: RiskAssessmentReport
    ):
        """Set document metadata properties."""
        core_props = doc.core_properties
        core_props.title = "Cybersecurity Risk Assessment Report"
        core_props.subject = "Risk Assessment"
        core_props.author = "Risk Assessment System"
        core_props.comments = f"Generated on {report_data.generated_at}"

    def _add_title_page(self, doc: Document, report_data: RiskAssessmentReport):
        """Add title page to document."""
        # Title
        title = doc.add_heading("Cybersecurity Risk Assessment Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Report ID and Date
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Report ID: {report_data.report_id}\n").bold = True
        info_para.add_run(
            f"Generated: {report_data.generated_at.strftime('%B %d, %Y %H:%M')}\n"
        )

        # Page break
        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, summary: ExecutiveSummary):
        """Add executive summary section."""
        doc.add_heading("Executive Summary", 1)

        # Summary statistics
        doc.add_paragraph(
            f"This report presents the findings of a comprehensive cybersecurity "
            f"risk assessment identifying {summary.total_vulnerabilities} "
            f"vulnerabilities requiring attention."
        )

        # Risk breakdown
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"

        risk_levels = [
            ("Critical Risk", summary.critical_risks, self.COLORS["critical"]),
            ("High Risk", summary.high_risks, self.COLORS["high"]),
            ("Medium Risk", summary.medium_risks, self.COLORS["medium"]),
            ("Low Risk", summary.low_risks, self.COLORS["low"]),
        ]

        table.rows[0].cells[0].text = "Risk Level"
        table.rows[0].cells[1].text = "Count"

        for i, (level, count, color) in enumerate(risk_levels, 1):
            cell0 = table.rows[i].cells[0]
            cell1 = table.rows[i].cells[1]

            cell0.text = level
            cell1.text = str(count)

            # Color the text
            run = cell0.paragraphs[0].runs[0]
            run.font.color.rgb = color
            run.font.bold = True

        doc.add_paragraph()

        # Key findings
        if summary.key_findings:
            doc.add_heading("Key Findings", 2)
            for finding in summary.key_findings:
                doc.add_paragraph(finding, style="List Bullet")

        # Top recommendations
        if summary.top_recommendations:
            doc.add_heading("Priority Recommendations", 2)
            for i, rec in enumerate(summary.top_recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style="List Number")

        doc.add_page_break()

    def _add_risk_heatmap(self, doc: Document, risk_ratings: List[RiskRating]):
        """Add risk heatmap visualization."""
        doc.add_heading("Risk Heatmap", 1)

        if not risk_ratings:
            doc.add_paragraph("No risk ratings available.")
            return

        # Create heatmap
        heatmap_path = self._create_risk_heatmap_chart(risk_ratings)

        if heatmap_path:
            doc.add_picture(heatmap_path, width=Inches(6))
            # Clean up temp file
            try:
                os.remove(heatmap_path)
            except:
                pass

        doc.add_paragraph()

    def _add_findings_overview(
        self, doc: Document, report_data: RiskAssessmentReport
    ):
        """Add findings overview table."""
        doc.add_heading("Findings Overview", 1)

        if not report_data.vulnerabilities:
            doc.add_paragraph("No vulnerabilities identified.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        headers = ["CVE ID", "Severity", "CVSS", "Exploited", "Risk Level"]

        for i, header in enumerate(headers):
            header_cells[i].text = header
            run = header_cells[i].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = self.COLORS["header"]

        # Data rows
        for vuln in report_data.vulnerabilities[:20]:  # Limit to top 20
            row_cells = table.add_row().cells
            cve = vuln.cve_detail

            row_cells[0].text = cve.cve_id
            row_cells[1].text = cve.cvss_severity or "N/A"
            row_cells[2].text = str(cve.cvss_score) if cve.cvss_score else "N/A"
            row_cells[3].text = "Yes" if vuln.exploitation_status.actively_exploited else "No"

            # Find corresponding risk rating
            risk_level = "N/A"
            for rating in report_data.risk_ratings:
                if rating.cve_id == cve.cve_id:
                    risk_level = rating.risk_level
                    break

            row_cells[4].text = risk_level

        doc.add_page_break()

    def _add_vulnerability_details(
        self, doc: Document, vulnerabilities: List[VulnerabilityAnalysis]
    ):
        """Add detailed vulnerability information."""
        doc.add_heading("Vulnerability Details", 1)

        for vuln in vulnerabilities[:10]:  # Limit to top 10 for detail
            cve = vuln.cve_detail

            # CVE heading
            doc.add_heading(cve.cve_id, 2)

            # Description
            doc.add_paragraph(f"Description: {cve.description[:500]}...")

            # Severity info
            severity_para = doc.add_paragraph()
            severity_para.add_run("Severity: ").bold = True
            severity_run = severity_para.add_run(
                f"{cve.cvss_severity or 'N/A'} ({cve.cvss_score or 'N/A'})"
            )

            if cve.cvss_severity == "CRITICAL":
                severity_run.font.color.rgb = self.COLORS["critical"]
            elif cve.cvss_severity == "HIGH":
                severity_run.font.color.rgb = self.COLORS["high"]

            # Exploitation status
            if vuln.exploitation_status.in_cisa_kev:
                exploit_para = doc.add_paragraph()
                exploit_run = exploit_para.add_run("⚠️  ACTIVELY EXPLOITED (CISA KEV)")
                exploit_run.font.color.rgb = self.COLORS["critical"]
                exploit_run.font.bold = True

            # Recommendation
            rec_para = doc.add_paragraph()
            rec_para.add_run("Recommendation: ").bold = True
            rec_para.add_run(vuln.recommendation)

            doc.add_paragraph()  # Spacer

        doc.add_page_break()

    def _add_threat_intelligence(self, doc: Document, threats: List[Any]):
        """Add threat intelligence section."""
        doc.add_heading("Threat Intelligence", 1)

        if not threats:
            doc.add_paragraph("No threat intelligence data available.")
            return

        for threat in threats[:5]:  # Limit to top 5
            doc.add_heading(f"Threat Analysis: {threat.cve_id}", 2)

            # MITRE ATT&CK techniques
            if threat.techniques:
                doc.add_paragraph("MITRE ATT&CK Techniques:", style="List Bullet")
                for tech in threat.techniques[:5]:
                    doc.add_paragraph(
                        f"{tech.technique_id}: {tech.name}", style="List Bullet 2"
                    )

            # Narrative
            if threat.narrative:
                doc.add_paragraph("Intelligence Summary:")
                doc.add_paragraph(threat.narrative[:500] + "...")

            # IOCs
            if threat.iocs:
                total_iocs = sum(len(v) for v in threat.iocs.values())
                if total_iocs > 0:
                    doc.add_paragraph(f"Indicators of Compromise: {total_iocs} identified")

            doc.add_paragraph()  # Spacer

        doc.add_page_break()

    def _add_risk_ratings(self, doc: Document, risk_ratings: List[RiskRating]):
        """Add risk ratings section."""
        doc.add_heading("Risk Ratings", 1)

        for rating in risk_ratings:
            doc.add_heading(f"{rating.cve_id} - {rating.asset_name}", 2)

            # Risk level
            risk_para = doc.add_paragraph()
            risk_para.add_run("Risk Level: ").bold = True
            risk_run = risk_para.add_run(
                f"{rating.risk_level} ({rating.risk_score}/25)"
            )

            color_key = rating.risk_level.lower()
            if color_key in self.COLORS:
                risk_run.font.color.rgb = self.COLORS[color_key]
            risk_run.font.bold = True

            # Likelihood and Impact
            doc.add_paragraph(
                f"Likelihood: {rating.likelihood.overall_score}/5 | "
                f"Impact: {rating.impact.overall_score}/5"
            )

            # Justification
            doc.add_paragraph("Justification:")
            doc.add_paragraph(rating.overall_justification)

            doc.add_paragraph()  # Spacer

        doc.add_page_break()

    def _add_recommendations(self, doc: Document, report_data: RiskAssessmentReport):
        """Add recommendations section."""
        doc.add_heading("Recommendations", 1)

        # Group by priority
        critical_recs = []
        high_recs = []
        medium_recs = []
        low_recs = []

        for rating in report_data.risk_ratings:
            rec_text = f"{rating.cve_id} ({rating.asset_name}): {rating.recommendations[0] if rating.recommendations else 'Review required'}"

            if rating.risk_level == "Critical":
                critical_recs.append(rec_text)
            elif rating.risk_level == "High":
                high_recs.append(rec_text)
            elif rating.risk_level == "Medium":
                medium_recs.append(rec_text)
            else:
                low_recs.append(rec_text)

        # Critical
        if critical_recs:
            heading = doc.add_heading("Critical Priority (Immediate Action)", 2)
            heading.runs[0].font.color.rgb = self.COLORS["critical"]
            for rec in critical_recs:
                doc.add_paragraph(rec, style="List Bullet")

        # High
        if high_recs:
            heading = doc.add_heading("High Priority (Within 1 Week)", 2)
            heading.runs[0].font.color.rgb = self.COLORS["high"]
            for rec in high_recs:
                doc.add_paragraph(rec, style="List Bullet")

        # Medium
        if medium_recs:
            heading = doc.add_heading("Medium Priority (Within 30 Days)", 2)
            heading.runs[0].font.color.rgb = self.COLORS["medium"]
            for rec in medium_recs:
                doc.add_paragraph(rec, style="List Bullet")

        # Low
        if low_recs:
            heading = doc.add_heading("Low Priority (Within 90 Days)", 2)
            heading.runs[0].font.color.rgb = self.COLORS["low"]
            for rec in low_recs:
                doc.add_paragraph(rec, style="List Bullet")

        doc.add_page_break()

    def _add_appendices(self, doc: Document, report_data: RiskAssessmentReport):
        """Add appendices section."""
        doc.add_heading("Appendices", 1)

        # Appendix A: Methodology
        doc.add_heading("Appendix A: Assessment Methodology", 2)
        doc.add_paragraph(
            "This risk assessment used a FAIR-based 5×5 risk matrix combining "
            "likelihood and impact dimensions:"
        )

        doc.add_paragraph("Likelihood Factors:", style="List Bullet")
        doc.add_paragraph("CVE Severity (CVSS Score)", style="List Bullet 2")
        doc.add_paragraph("Exploitation Status (CISA KEV, VirusTotal)", style="List Bullet 2")
        doc.add_paragraph("Asset Exposure", style="List Bullet 2")
        doc.add_paragraph("Threat Actor Capability", style="List Bullet 2")
        doc.add_paragraph("Control Effectiveness", style="List Bullet 2")

        doc.add_paragraph("Impact Factors:", style="List Bullet")
        doc.add_paragraph("Asset Criticality", style="List Bullet 2")
        doc.add_paragraph("Data Sensitivity", style="List Bullet 2")
        doc.add_paragraph("Business Impact", style="List Bullet 2")
        doc.add_paragraph("Compliance Impact", style="List Bullet 2")
        doc.add_paragraph("Operational Impact", style="List Bullet 2")

    def _create_risk_heatmap_chart(self, risk_ratings: List[RiskRating]) -> Optional[str]:
        """Create risk heatmap chart using matplotlib.

        Args:
            risk_ratings: List of risk ratings

        Returns:
            Path to saved chart image or None
        """
        try:
            # Count risks by level
            risk_counts = {
                "Critical": 0,
                "High": 0,
                "Medium": 0,
                "Low": 0,
            }

            for rating in risk_ratings:
                risk_counts[rating.risk_level] += 1

            # Create figure
            fig, ax = plt.subplots(figsize=(10, 6))

            # Data
            levels = list(risk_counts.keys())
            counts = list(risk_counts.values())
            colors = ["#DC143C", "#FF6347", "#FFA500", "#228B22"]

            # Create bar chart
            bars = ax.bar(levels, counts, color=colors, edgecolor="black", linewidth=1.5)

            # Customize
            ax.set_xlabel("Risk Level", fontsize=12, fontweight="bold")
            ax.set_ylabel("Count", fontsize=12, fontweight="bold")
            ax.set_title("Risk Distribution", fontsize=14, fontweight="bold")
            ax.grid(axis="y", alpha=0.3)

            # Add value labels on bars
            for bar in bars:
                height = bar.get_height()
                ax.text(
                    bar.get_x() + bar.get_width() / 2.0,
                    height,
                    f"{int(height)}",
                    ha="center",
                    va="bottom",
                    fontweight="bold",
                )

            plt.tight_layout()

            # Save to temp file
            temp_path = os.path.join(self.output_dir, "temp_heatmap.png")
            plt.savefig(temp_path, dpi=300, bbox_inches="tight")
            plt.close()

            return temp_path

        except Exception as e:
            logger.error(f"Error creating heatmap: {e}")
            return None
