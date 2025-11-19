"""Document parsing utilities for PDF, DOCX, and XLSX files.

This module provides functions to parse various document formats and extract
text content, metadata, and structured data for security analysis.
"""

import sys
from pathlib import Path

# Ensure src is in path for absolute imports
_src_path = str(Path(__file__).parent.parent)
if _src_path not in sys.path:
    sys.path.insert(0, _src_path)

import os
import re
from typing import Optional, Dict, Any, List
from pathlib import Path
import logging

# PDF parsing
from pypdf import PdfReader

# DOCX parsing
from docx import Document

# Excel parsing
import openpyxl
import pandas as pd

from models.schemas import DocumentMetadata, ExtractedEntity, DocumentAnalysis

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for various document formats.

    Supports PDF, DOCX, XLSX, TXT, MD, and CSV files with metadata extraction
    and text content parsing.

    Week 7 Enhancement: Added .txt, .md, .csv support and document classification.
    """

    SUPPORTED_FORMATS = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".md", ".csv", ".pptx"}

    def __init__(self):
        """Initialize document parser."""
        logger.info("Document parser initialized")

    def parse_document(self, file_path: str) -> Optional[DocumentAnalysis]:
        """Parse a document and return analysis.

        Args:
            file_path: Path to document file

        Returns:
            DocumentAnalysis object or None if parsing fails

        Example:
            >>> parser = DocumentParser()
            >>> analysis = parser.parse_document("report.pdf")
            >>> if analysis:
            ...     print(f"Text length: {len(analysis.text_content)}")
            ...     print(f"Entities found: {len(analysis.entities)}")
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None

        file_ext = Path(file_path).suffix.lower()

        if file_ext not in self.SUPPORTED_FORMATS:
            logger.error(f"Unsupported file format: {file_ext}")
            return None

        try:
            # Parse based on file type
            if file_ext == ".pdf":
                return self._parse_pdf(file_path)
            elif file_ext == ".docx":
                return self._parse_docx(file_path)
            elif file_ext in {".xlsx", ".xls"}:
                return self._parse_excel(file_path)
            elif file_ext == ".pptx":
                return self.parse_pptx(file_path)
            elif file_ext == ".txt":
                return self._parse_txt(file_path)
            elif file_ext == ".md":
                return self._parse_markdown(file_path)
            elif file_ext == ".csv":
                return self._parse_csv(file_path)

        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            return None

    def _parse_pdf(self, file_path: str) -> DocumentAnalysis:
        """Parse PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            DocumentAnalysis object
        """
        logger.info(f"Parsing PDF: {file_path}")

        reader = PdfReader(file_path)

        # Extract metadata
        metadata_dict = reader.metadata or {}
        metadata = DocumentMetadata(
            filename=os.path.basename(file_path),
            file_type="pdf",
            page_count=len(reader.pages),
            author=metadata_dict.get("/Author"),
            created_date=str(metadata_dict.get("/CreationDate", "")),
            modified_date=str(metadata_dict.get("/ModDate", "")),
        )

        # Extract text from all pages
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Error extracting text from page {page_num}: {e}")

        full_text = "\n".join(text_parts)

        # Extract entities
        entities = self._extract_entities(full_text)

        # Generate summary
        summary = self._generate_summary(full_text, len(reader.pages))

        return DocumentAnalysis(
            metadata=metadata,
            text_content=full_text,
            entities=entities,
            summary=summary,
        )

    def _parse_docx(self, file_path: str) -> DocumentAnalysis:
        """Parse DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            DocumentAnalysis object
        """
        logger.info(f"Parsing DOCX: {file_path}")

        doc = Document(file_path)

        # Extract metadata from core properties
        core_props = doc.core_properties
        metadata = DocumentMetadata(
            filename=os.path.basename(file_path),
            file_type="docx",
            page_count=None,  # DOCX doesn't have reliable page count
            author=core_props.author,
            created_date=str(core_props.created) if core_props.created else None,
            modified_date=str(core_props.modified) if core_props.modified else None,
        )

        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        # Extract entities
        entities = self._extract_entities(full_text)

        # Generate summary
        summary = self._generate_summary(full_text, len(doc.paragraphs))

        return DocumentAnalysis(
            metadata=metadata,
            text_content=full_text,
            entities=entities,
            summary=summary,
        )

    def _parse_excel(self, file_path: str) -> DocumentAnalysis:
        """Parse Excel file.

        Args:
            file_path: Path to Excel file

        Returns:
            DocumentAnalysis object
        """
        logger.info(f"Parsing Excel: {file_path}")

        # Load workbook for metadata
        wb = openpyxl.load_workbook(file_path, data_only=True)

        metadata = DocumentMetadata(
            filename=os.path.basename(file_path),
            file_type="xlsx",
            page_count=len(wb.sheetnames),
            author=wb.properties.creator,
            created_date=str(wb.properties.created) if wb.properties.created else None,
            modified_date=str(wb.properties.modified) if wb.properties.modified else None,
        )

        # Extract text from all sheets using pandas
        text_parts = []

        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                # Add sheet name
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")

                # Convert dataframe to text
                # Include headers and data
                text_parts.append(df.to_string())

        except Exception as e:
            logger.warning(f"Error reading Excel sheets with pandas: {e}")

            # Fallback to openpyxl
            for sheet in wb.worksheets:
                text_parts.append(f"\n=== Sheet: {sheet.title} ===\n")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        # Extract entities
        entities = self._extract_entities(full_text)

        # Generate summary
        summary = self._generate_summary(full_text, len(wb.sheetnames))

        wb.close()

        return DocumentAnalysis(
            metadata=metadata,
            text_content=full_text,
            entities=entities,
            summary=summary,
        )

    def _extract_entities(self, text: str) -> List[ExtractedEntity]:
        """Extract security-related entities from text.

        Uses regex patterns to identify CVEs, controls, assets, risks, and findings.

        Args:
            text: Text to analyze

        Returns:
            List of ExtractedEntity objects
        """
        entities = []

        # CVE pattern
        cve_pattern = r"CVE-\d{4}-\d{4,7}"
        for match in re.finditer(cve_pattern, text, re.IGNORECASE):
            cve_id = match.group().upper()
            start = max(0, match.start() - 50)
            end = min(len(text), match.end() + 50)
            context = text[start:end].replace("\n", " ")

            entities.append(
                ExtractedEntity(
                    entity_type="cve",
                    value=cve_id,
                    confidence=1.0,
                    context=context,
                )
            )

        # Control identifiers (e.g., NIST controls, ISO controls)
        control_patterns = [
            r"\b(AC|AT|AU|CA|CM|CP|IA|IR|MA|MP|PE|PL|PS|RA|SA|SC|SI)-\d+",  # NIST
            r"\bISO[\s-]?\d{5}",  # ISO
            r"\b(CIS|PCI|SOX|HIPAA|GDPR)[\s-]?\d+",  # Various frameworks
        ]

        for pattern in control_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                control_id = match.group()
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].replace("\n", " ")

                entities.append(
                    ExtractedEntity(
                        entity_type="control",
                        value=control_id,
                        confidence=0.9,
                        context=context,
                    )
                )

        # Asset patterns (servers, databases, etc.)
        asset_patterns = [
            r"\b(server|database|web[-\s]?server|app[-\s]?server|db|host)[-\s]?\w+",
            r"\b\w+[-\s]?(prod|production|dev|development|test|staging)\b",
        ]

        for pattern in asset_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                asset_name = match.group()
                start = max(0, match.start() - 30)
                end = min(len(text), match.end() + 30)
                context = text[start:end].replace("\n", " ")

                entities.append(
                    ExtractedEntity(
                        entity_type="asset",
                        value=asset_name,
                        confidence=0.7,
                        context=context,
                    )
                )

        # Risk keywords
        risk_keywords = [
            "critical risk",
            "high risk",
            "medium risk",
            "low risk",
            "vulnerability",
            "threat",
            "weakness",
        ]

        for keyword in risk_keywords:
            for match in re.finditer(re.escape(keyword), text, re.IGNORECASE):
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 100)
                context = text[start:end].replace("\n", " ")

                entities.append(
                    ExtractedEntity(
                        entity_type="risk",
                        value=match.group(),
                        confidence=0.8,
                        context=context,
                    )
                )

        # Finding keywords
        finding_keywords = [
            "finding",
            "observation",
            "recommendation",
            "remediation",
            "mitigation",
        ]

        for keyword in finding_keywords:
            for match in re.finditer(re.escape(keyword), text, re.IGNORECASE):
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 100)
                context = text[start:end].replace("\n", " ")

                entities.append(
                    ExtractedEntity(
                        entity_type="finding",
                        value=match.group(),
                        confidence=0.7,
                        context=context,
                    )
                )

        # Remove duplicate entities
        unique_entities = []
        seen = set()

        for entity in entities:
            key = (entity.entity_type, entity.value.lower())
            if key not in seen:
                seen.add(key)
                unique_entities.append(entity)

        logger.info(f"Extracted {len(unique_entities)} entities from document")
        return unique_entities

    def _generate_summary(self, text: str, element_count: int) -> str:
        """Generate a summary of the document.

        Args:
            text: Full document text
            element_count: Number of pages/paragraphs/sheets

        Returns:
            Summary string
        """
        word_count = len(text.split())

        # Count entity types
        cves = len(re.findall(r"CVE-\d{4}-\d{4,7}", text, re.IGNORECASE))
        risks = len(re.findall(r"\b(critical|high|medium|low)\s+risk\b", text, re.IGNORECASE))

        summary = (
            f"Document contains {word_count:,} words across {element_count} sections. "
            f"Identified {cves} CVE references and {risks} risk statements."
        )

        return summary

    def extract_cves_from_document(self, file_path: str) -> List[str]:
        """Extract just CVE IDs from a document.

        Args:
            file_path: Path to document

        Returns:
            List of unique CVE IDs

        Example:
            >>> parser = DocumentParser()
            >>> cves = parser.extract_cves_from_document("vuln_report.pdf")
            >>> print(f"Found CVEs: {cves}")
        """
        analysis = self.parse_document(file_path)

        if not analysis:
            return []

        cves = [
            entity.value
            for entity in analysis.entities
            if entity.entity_type == "cve"
        ]

        return list(set(cves))  # Return unique CVEs

    def _parse_txt(self, file_path: str) -> DocumentAnalysis:
        """Parse plain text file (Week 7 enhancement).

        Args:
            file_path: Path to TXT file

        Returns:
            DocumentAnalysis object
        """
        logger.info(f"Parsing TXT: {file_path}")

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text_content = f.read()

        # Get file stats
        file_stats = os.stat(file_path)

        metadata = DocumentMetadata(
            filename=os.path.basename(file_path),
            file_type="txt",
            page_count=1,  # Text files considered as single page
            file_size=file_stats.st_size,
            created_date=str(file_stats.st_mtime)
        )

        # Extract entities
        entities = self._extract_entities(text_content)

        # Classify document
        doc_classification = self._classify_document(text_content)
        tags = self._generate_tags(text_content)

        return DocumentAnalysis(
            metadata=metadata,
            text_content=text_content,
            entities=entities,
            summary=f"Text document with {len(text_content)} characters",
            confidence_score=0.95,  # High confidence for plain text
            document_classification=doc_classification,
            automated_tags=tags
        )

    def _parse_markdown(self, file_path: str) -> DocumentAnalysis:
        """Parse markdown file (Week 7 enhancement).

        Args:
            file_path: Path to MD file

        Returns:
            DocumentAnalysis object
        """
        logger.info(f"Parsing Markdown: {file_path}")

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text_content = f.read()

        # Get file stats
        file_stats = os.stat(file_path)

        # Count markdown sections (headers)
        section_count = len(re.findall(r'^#{1,6}\s', text_content, re.MULTILINE))

        metadata = DocumentMetadata(
            filename=os.path.basename(file_path),
            file_type="markdown",
            page_count=max(1, section_count),  # Use section count as page equivalent
            file_size=file_stats.st_size,
            created_date=str(file_stats.st_mtime)
        )

        # Extract entities
        entities = self._extract_entities(text_content)

        # Classify document
        doc_classification = self._classify_document(text_content)
        tags = self._generate_tags(text_content)

        return DocumentAnalysis(
            metadata=metadata,
            text_content=text_content,
            entities=entities,
            summary=f"Markdown document with {section_count} sections",
            confidence_score=0.90,
            document_classification=doc_classification,
            automated_tags=tags
        )

    def _parse_csv(self, file_path: str) -> DocumentAnalysis:
        """Parse CSV file (Week 7 enhancement).

        Args:
            file_path: Path to CSV file

        Returns:
            DocumentAnalysis object
        """
        logger.info(f"Parsing CSV: {file_path}")

        # Read CSV into DataFrame
        df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')

        # Convert to text representation
        text_content = df.to_string()

        # Also get column info
        column_info = f"Columns: {', '.join(df.columns.tolist())}\n\n"
        text_content = column_info + text_content

        # Get file stats
        file_stats = os.stat(file_path)

        metadata = DocumentMetadata(
            filename=os.path.basename(file_path),
            file_type="csv",
            page_count=1,  # CSV considered as single page
            file_size=file_stats.st_size,
            created_date=str(file_stats.st_mtime)
        )

        # Extract entities
        entities = self._extract_entities(text_content)

        # Classify document
        doc_classification = self._classify_document(text_content)
        tags = self._generate_tags(text_content)
        tags.append(f"rows:{len(df)}")  # Add row count as tag
        tags.append(f"cols:{len(df.columns)}")  # Add column count as tag

        return DocumentAnalysis(
            metadata=metadata,
            text_content=text_content,
            entities=entities,
            summary=f"CSV data with {len(df)} rows and {len(df.columns)} columns",
            confidence_score=0.95,
            document_classification=doc_classification,
            automated_tags=tags
        )

    def _classify_document(self, text: str) -> str:
        """Classify document type based on content (Week 7 enhancement).

        Args:
            text: Document text content

        Returns:
            Document classification string
        """
        text_lower = text.lower()

        # Security-related keywords
        security_keywords = ['vulnerability', 'threat', 'attack', 'malware', 'exploit',
                           'cve-', 'security', 'breach', 'intrusion']
        risk_keywords = ['risk', 'assessment', 'evaluation', 'mitigation', 'impact']
        audit_keywords = ['audit', 'compliance', 'regulatory', 'framework', 'standard']
        policy_keywords = ['policy', 'procedure', 'guideline', 'requirement']
        incident_keywords = ['incident', 'response', 'forensic', 'investigation']

        # Count keyword occurrences
        security_count = sum(1 for kw in security_keywords if kw in text_lower)
        risk_count = sum(1 for kw in risk_keywords if kw in text_lower)
        audit_count = sum(1 for kw in audit_keywords if kw in text_lower)
        policy_count = sum(1 for kw in policy_keywords if kw in text_lower)
        incident_count = sum(1 for kw in incident_keywords if kw in text_lower)

        # Classify based on highest count
        counts = {
            'security_report': security_count,
            'risk_assessment': risk_count,
            'audit_report': audit_count,
            'policy_document': policy_count,
            'incident_report': incident_count
        }

        classification = max(counts, key=counts.get)

        # If no strong match, classify as general
        if counts[classification] < 2:
            classification = 'general_document'

        return classification

    def _generate_tags(self, text: str) -> List[str]:
        """Generate automated tags based on content (Week 7 enhancement).

        Args:
            text: Document text content

        Returns:
            List of tags
        """
        tags = []
        text_lower = text.lower()

        # Domain tags
        if any(kw in text_lower for kw in ['vulnerability', 'exploit', 'cve']):
            tags.append('vulnerability-management')
        if any(kw in text_lower for kw in ['firewall', 'network', 'ids', 'ips']):
            tags.append('network-security')
        if any(kw in text_lower for kw in ['encryption', 'crypto', 'cipher']):
            tags.append('cryptography')
        if any(kw in text_lower for kw in ['authentication', 'authorization', 'access-control']):
            tags.append('access-management')
        if any(kw in text_lower for kw in ['compliance', 'regulatory', 'gdpr', 'hipaa']):
            tags.append('compliance')
        if any(kw in text_lower for kw in ['incident', 'breach', 'forensic']):
            tags.append('incident-response')

        return tags

    # ========== Week 7 Session 2: Document Intelligence Features ==========

    def parse_scanned_pdf(self, pdf_path: str) -> Optional[DocumentAnalysis]:
        """
        Parse scanned PDF using OCR.

        Args:
            pdf_path: Path to scanned PDF file

        Returns:
            DocumentAnalysis object with OCR-extracted text
        """
        try:
            from .ocr_processor import OCRProcessor

            logger.info(f"Processing scanned PDF with OCR: {pdf_path}")

            # Initialize OCR processor
            ocr = OCRProcessor()

            # Check if PDF is scanned
            if not ocr.is_scanned_pdf(pdf_path):
                logger.info("PDF appears to be native (not scanned), using standard parser")
                return self._parse_pdf(pdf_path)

            # Extract text from scanned PDF
            ocr_results = ocr.extract_text_from_scanned_pdf(pdf_path)

            if not ocr_results:
                logger.warning("OCR extraction returned no results")
                return None

            # Combine text from all pages
            all_text = "\n\n".join(
                result['text'] for result in ocr_results if result.get('text')
            )

            # Calculate average confidence
            avg_confidence = sum(r['confidence'] for r in ocr_results) / len(ocr_results)

            # Create metadata
            metadata = DocumentMetadata(
                file_path=pdf_path,
                file_name=Path(pdf_path).name,
                file_size=os.path.getsize(pdf_path),
                page_count=len(ocr_results),
                doc_type='pdf_scanned',
                extraction_method='ocr',
                ocr_confidence=avg_confidence
            )

            # Extract entities
            entities = self._extract_security_entities(all_text)

            # Generate tags
            tags = self._generate_security_tags(all_text)

            return DocumentAnalysis(
                metadata=metadata,
                text_content=all_text,
                entities=entities,
                tags=tags
            )

        except ImportError:
            logger.error("OCR processor not available")
            return None
        except Exception as e:
            logger.error(f"Error processing scanned PDF: {e}")
            return None

    def extract_tables(self, file_path: str) -> List[Any]:
        """
        Extract tables from document.

        Args:
            file_path: Path to document file

        Returns:
            List of DataFrames representing tables
        """
        try:
            from .table_extractor import TableExtractor

            file_ext = Path(file_path).suffix.lower()

            if file_ext == '.pdf':
                extractor = TableExtractor()
                tables = extractor.extract_tables_from_pdf(file_path)
                logger.info(f"Extracted {len(tables)} tables from PDF")
                return tables
            elif file_ext == '.pptx':
                from .pptx_parser import PPTXParser
                parser = PPTXParser(extract_tables=True)
                tables = parser.extract_all_tables(file_path)
                logger.info(f"Extracted {len(tables)} tables from PowerPoint")
                return tables
            else:
                logger.warning(f"Table extraction not supported for {file_ext}")
                return []

        except ImportError as e:
            logger.error(f"Table extraction dependencies not available: {e}")
            return []
        except Exception as e:
            logger.error(f"Error extracting tables: {e}")
            return []

    def classify_document_type(self, text: str) -> tuple:
        """
        Classify document type using ML.

        Args:
            text: Document text

        Returns:
            Tuple of (document_type, confidence_score)
        """
        try:
            from .document_classifier import DocumentClassifier

            classifier = DocumentClassifier()
            doc_type, confidence = classifier.predict_with_confidence(text)

            logger.info(f"Document classified as: {doc_type} (confidence: {confidence:.2f})")

            return (doc_type, confidence)

        except ImportError:
            logger.error("Document classifier not available")
            return ('unknown', 0.0)
        except Exception as e:
            logger.error(f"Error classifying document: {e}")
            return ('unknown', 0.0)

    def parse_pptx(self, pptx_path: str) -> Optional[DocumentAnalysis]:
        """
        Parse PowerPoint presentation.

        Args:
            pptx_path: Path to PPTX file

        Returns:
            DocumentAnalysis object
        """
        try:
            from .pptx_parser import PPTXParser

            logger.info(f"Parsing PowerPoint: {pptx_path}")

            parser = PPTXParser()
            result = parser.parse_presentation(pptx_path)

            if not result['success']:
                logger.error(f"Failed to parse PowerPoint: {result.get('error')}")
                return None

            # Get combined text
            summary = result.get('summary', {})
            all_text = summary.get('all_text', '')
            all_notes = summary.get('all_notes', '')
            combined_text = f"{all_text}\n\n--- Speaker Notes ---\n\n{all_notes}" if all_notes else all_text

            # Create metadata
            pptx_metadata = result.get('metadata', {})
            metadata = DocumentMetadata(
                file_path=pptx_path,
                file_name=Path(pptx_path).name,
                file_size=os.path.getsize(pptx_path),
                page_count=summary.get('total_slides', 0),
                doc_type='pptx',
                author=pptx_metadata.get('author', ''),
                title=pptx_metadata.get('title', ''),
                created_date=pptx_metadata.get('created'),
                modified_date=pptx_metadata.get('modified')
            )

            # Extract entities
            entities = self._extract_security_entities(combined_text)

            # Generate tags
            tags = self._generate_security_tags(combined_text)

            return DocumentAnalysis(
                metadata=metadata,
                text_content=combined_text,
                entities=entities,
                tags=tags
            )

        except ImportError:
            logger.error("PowerPoint parser not available")
            return None
        except Exception as e:
            logger.error(f"Error parsing PowerPoint: {e}")
            return None

    def auto_detect_format(self, file_path: str) -> str:
        """
        Auto-detect document format and processing requirements.

        Args:
            file_path: Path to document file

        Returns:
            Detected format/type ('native_pdf', 'scanned_pdf', 'pptx', etc.)
        """
        try:
            file_ext = Path(file_path).suffix.lower()

            if file_ext == '.pdf':
                # Check if scanned
                from .ocr_processor import OCRProcessor
                ocr = OCRProcessor()
                if ocr.is_scanned_pdf(file_path):
                    return 'scanned_pdf'
                else:
                    return 'native_pdf'
            elif file_ext == '.pptx':
                return 'pptx'
            elif file_ext == '.docx':
                return 'docx'
            elif file_ext in {'.xlsx', '.xls'}:
                return 'excel'
            elif file_ext == '.txt':
                return 'text'
            elif file_ext == '.md':
                return 'markdown'
            elif file_ext == '.csv':
                return 'csv'
            else:
                return 'unknown'

        except Exception as e:
            logger.error(f"Error auto-detecting format: {e}")
            return 'unknown'
